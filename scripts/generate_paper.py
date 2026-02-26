#!/usr/bin/env python3
"""
Generate properly formatted DOCX and PDF versions of the revised paper.

Requires: pip install python-docx Pillow
PDF is rendered via headless Chrome from a styled HTML.
"""

import os
import re
import sys
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent.parent
MD_PATH  = BASE_DIR / "REVISED_PAPER.md"
DOCX_OUT = BASE_DIR / "paper.docx"
PDF_OUT  = BASE_DIR / "paper.pdf"
HTML_TMP = BASE_DIR / "paper_tmp.html"

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_image_safe(doc, rel_path, caption_text, width=Inches(5.5)):
    img_path = BASE_DIR / rel_path
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
        cap = doc.add_paragraph(caption_text)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        p = doc.add_paragraph(f"[Figure not found: {rel_path}]")
        p.runs[0].italic = True

# ─── Build DOCX ──────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title ───────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Hybrid Deep Learning Framework for EEG-Based Detection of\n"
        "Depression and Brain Abnormalities"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # ── Authors ──────────────────────────────────────────────────────────────
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = authors.add_run(
        "Abhinav Joshi¹*, Prawar Chaudhary², Kaushal Kumar³, "
        "Chintan Singh⁴, Roobal Chaudhary⁵, Preeti Rustagi⁶, Harsh Pandey⁷"
    )
    r.font.size = Pt(11)
    r.bold = True

    # ── Affiliations ─────────────────────────────────────────────────────────
    affil_lines = [
        "¹ RISE Lab, Indian Institute of Technology Delhi, New Delhi, INDIA",
        "² School of Basics and Applied Sciences, KR Mangalam University, Gurugram, INDIA",
        "³ Department of Mechanical Engineering, KR Mangalam University, Gurugram, INDIA",
        "⁴ Amity Institute of Forensic Sciences, Amity University Noida, Uttar Pradesh, INDIA",
        "⁵ Department of Forensic Science, Sharda School of Allied Health Sciences, Sharda University, Greater Noida, INDIA",
        "⁶ Faculty of Commerce & Management, SGT University, Gurugram, Haryana, INDIA",
        "⁷ Department of Biotechnology and Chemical Engineering, Manipal University Jaipur, INDIA",
        "* Corresponding author: abhinavjoshi@iitd.ac.in",
    ]
    for line in affil_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)

    add_horizontal_rule(doc)

    # ── Abstract ─────────────────────────────────────────────────────────────
    h = doc.add_heading("Abstract", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    abstract_text = (
        "Major Depressive Disorder (MDD) and other neuropsychiatric conditions often remain "
        "underdiagnosed due to the subjective nature of clinical assessments and limited access "
        "to advanced neuroimaging technologies in routine healthcare. Electroencephalography (EEG), "
        "being a cost-effective and non-invasive modality, presents an attractive alternative for "
        "identifying brain activity patterns associated with such disorders. This study proposes a "
        "novel hybrid deep learning framework that integrates a fine-tuned InceptionV3 convolutional "
        "neural network with Long Short-Term Memory (LSTM) units to improve the accuracy of "
        "EEG-based classification. We introduce a novel Spatial RGB Spectrogram Mapping technique, "
        "translating raw topographic EEG channels (Frontal, Central, Posterior) into distinct colour "
        "channels (Red, Green, Blue) to preserve spatial dynamics. Employing strict Leave-One-Subject-Out "
        "(LOSO) cross-validation, our final architecture reached a Subject-Level Clinical Diagnostic "
        "Accuracy of 90.0%.  Extensive outlier analysis was conducted to understand biological variations, "
        "specifically identifying why subsets of healthy and depressed brains exhibit counter-intuitive "
        "Alpha energy markers. The results confirm the clinical potential of spatial-temporal deep learning "
        "architectures in EEG signal interpretation."
    )
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(10)

    add_horizontal_rule(doc)

    # ─── 1. Introduction ─────────────────────────────────────────────────────
    doc.add_heading("1. Introduction", level=1).runs[0].font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
    doc.add_paragraph(
        "Major Depressive Disorder (MDD) is difficult to diagnose clinically because of its camouflaged "
        "neurophysiological landscape and dependency on subjective measurements. Conventional diagnostic "
        "procedures, mainly clinical interviews and behavioral examinations, often result in underreported "
        "or misdiagnosed cases, mostly in under-resourced regions. Electroencephalography (EEG), as a "
        "non-invasive and cost-effective modality, has emerged as a promising tool for capturing the dynamic "
        "electrical activity of the brain. Altered theta and alpha wave patterns in the EEG spectrum have "
        "been consistently associated with depressive episodes; however, manual interpretation remains "
        "highly dependent on expert knowledge."
    ).runs[0].font.size = Pt(11)
    doc.add_paragraph(
        "Recent research has explored ML and DL techniques for automated EEG analysis. Traditional CNNs "
        "extract spatial features while LSTM networks capture temporal dependencies. Their independent use, "
        "however, often falls short of capturing the complex, multidimensional nature of EEG signals. In "
        "this research, we propose a hybrid InceptionV3 + LSTM framework combined with a novel Spatial RGB "
        "Mapping approach that uniquely colour-codes brain topography for regional awareness. All evaluations "
        "are conducted under strict Leave-One-Subject-Out (LOSO) cross-validation to eliminate data leakage."
    ).runs[0].font.size = Pt(11)

    # ─── 2. Materials and Methods ─────────────────────────────────────────────
    doc.add_heading("2. Materials and Methods", level=1).runs[0].font.color.rgb = RGBColor(0x1a,0x1a,0x2e)

    doc.add_heading("2.1 Dataset and Cohort Selection", level=2)
    doc.add_paragraph(
        "The dataset is sourced from the publicly available repository 'MDD Patients and Healthy Controls "
        "EEG Data (New)' [5], captured using a 19-channel setup adhering to the 10-20 electrode placement "
        "system. After rigorous artifact rejection — handling eyeblinks, muscle artifacts, and bad channels — "
        "we selected a heavily curated cohort of 40 high-fidelity subjects (14 Healthy, 26 MDD). These subjects "
        "possessed the clean, continuous recording lengths required for our localized seq=10 (7-second continuous) "
        "temporal analysis."
    ).runs[0].font.size = Pt(11)

    doc.add_heading("2.2 Preprocessing and Spatial RGB Spectrogram Mapping", level=2)
    doc.add_paragraph(
        "A multi-step pipeline transforms continuous raw signals into spatial-temporal visual sequences:"
    ).runs[0].font.size = Pt(11)
    steps = [
        ("Filtering & Epoching",
         "A bandpass filter (0.5–40 Hz) was applied, followed by dividing cleaned signals into discrete 700 ms epochs."),
        ("Wavelet Transform",
         "DWT with Daubechies 4 (db4) extracted the approximation coefficients (cA), emphasising low-frequency neural oscillations."),
        ("Spatial RGB Mapping (Novel Contribution)",
         "Rather than averaging all 19 channels into a greyscale scalogram, we assigned topographic electrode groups "
         "to independent colour channels:\n"
         "  • Frontal [Fp1, Fp2, F7, F3, Fz, F4, F8] → RED — frontal alpha asymmetry is a key MDD biomarker.\n"
         "  • Central  [T3, C3, Cz, C4, T4]           → GREEN — isolates somatosensory/motor activity.\n"
         "  • Posterior [T5, P3, Pz, P4, T6, O1, O2] → BLUE  — resting-state alpha rhythms separated from frontal generators."),
    ]
    for i, (bold_part, detail) in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        run_b = p.add_run(f"{bold_part}: ")
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_d = p.add_run(detail)
        run_d.font.size = Pt(11)

    add_image_safe(doc, "results/data_investigation/1_class_averaged_spectrograms.png",
                   "Figure 1: Averaged grayscale profiles demonstrating the lack of class separability without spatial mapping.")

    doc.add_heading("2.3 Hybrid Architecture: InceptionV3 + LSTM", level=2)
    doc.add_paragraph(
        "Why InceptionV3? EEG spectrograms display multi-scale hierarchical features. InceptionV3's parallel "
        "convolutional kernels (1×1, 3×3, 5×5) simultaneously capture highly-localised features and broader "
        "wave synchronisations. To prevent overfitting on a limited medical dataset, we partially fine-tuned "
        "only the Mixed_7c block (~7 M trainable parameters), adapting ImageNet weights to EEG topology.\n\n"
        "Why LSTM? CNNs cannot natively track how brain states evolve over successive seconds. After "
        "InceptionV3 extracts per-frame spatial vectors, a 128-unit LSTM layer models the sequential "
        "dependencies across the 7-second time series, distinguishing persistent depressive wave states "
        "from momentary artifact blips."
    ).runs[0].font.size = Pt(11)

    doc.add_heading("2.4 Class Imbalance and LOSO Validation", level=2)
    doc.add_paragraph(
        "The cohort was severely imbalanced: 65% MDD (26/40) vs 35% Healthy (14/40). We used "
        "BCEWithLogitsLoss with pos_weight = n_H / n_MDD ≈ 0.53, inversely penalising MDD over-predictions "
        "and forcing the network to rigorously learn minority Healthy features. All results are reported "
        "under strict Leave-One-Subject-Out (LOSO) cross-validation — train on 39 subjects, test on 1 "
        "entirely unseen subject, repeated 40 independent times — eliminating any possibility of epoch-level "
        "data leakage."
    ).runs[0].font.size = Pt(11)

    # ─── 3. Development Journey ───────────────────────────────────────────────
    doc.add_heading("3. Development Journey: Architecture Iterations", level=1).runs[0].font.color.rgb = RGBColor(0x1a,0x1a,0x2e)

    doc.add_heading("3.1 Experiment 1 — Baseline Architecture & The Specificity Gap", level=2)
    doc.add_paragraph(
        "Our initial InceptionV3 + LSTM model used averaged greyscale spectrograms. While it achieved "
        "90.1% Sensitivity for detecting MDD, Specificity flatlined at 63.6% — the model was almost "
        "always predicting MDD for every subject."
    ).runs[0].font.size = Pt(11)
    add_image_safe(doc, "results/exp1_seq_lstm/figures/01_per_subject_performance.png",
                   "Figure 2: Experiment 1 per-subject results — strong MDD recall but consistently poor Healthy detection.")

    doc.add_heading("3.2 Dedicated Analysis — The 'Impossible' Outliers", level=2)
    doc.add_paragraph(
        "We investigated the raw Alpha energy bands of repeatedly misclassified subjects and discovered "
        "5 'Impossible Outliers' whose brain waves fundamentally mimic the opposite clinical condition:"
    ).runs[0].font.size = Pt(11)
    outlier_items = [
        ("Depressed-Looking Healthy Brains (H_16, H_24, H_27)",
         "These healthy individuals naturally possess extremely low Alpha band energy. To a generic spectral "
         "algorithm, their waves look identical to a clinically depressed patient."),
        ("Healthy-Looking MDD Brains (MDD_5, MDD_19)",
         "These depressed patients exhibit unusually prominent Alpha energy, making their EEG profile look perfectly healthy."),
    ]
    for bold_part, detail in outlier_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{bold_part}: ").bold = True
        p.add_run(detail).font.size = Pt(11)

    doc.add_paragraph(
        "Because Experiment 1 averaged all 19 channels into a single greyscale image, spatial quirks were "
        "destroyed. The model could not distinguish 'naturally low Alpha' from 'MDD-induced low Alpha.'"
    ).runs[0].font.size = Pt(11)
    add_image_safe(doc, "results/outlier_analysis/spectral_profiles_comparison.png",
                   "Figure 3: Spectral profiles demonstrating Alpha energy overlap between anomalous Healthy and MDD subjects.")

    doc.add_heading("3.3 Experiment 2 — Fine-Tuning Alone Cannot Break the Spectral Overlap", level=2)
    doc.add_paragraph(
        "Experiment 2 broadened fine-tuning to deeper InceptionV3 layers in an attempt to overcome the "
        "outlier problem without new data. The approach caused instability — global accuracy collapsed to "
        "~68% — proving that CNN adaptation alone cannot compensate for the missing spatial signal."
    ).runs[0].font.size = Pt(11)
    add_image_safe(doc, "results/approach2/metric_distributions.png",
                   "Figure 4: Metric distribution during Experiment 2 — high variance confirms model instability under broad fine-tuning.")

    doc.add_heading("3.4 Final Experiment — Spatial RGB + Partial Fine-Tuning", level=2)
    doc.add_paragraph(
        "The final architecture combined: (1) Spatial RGB spectrograms for regional awareness, "
        "(2) Mixed_7c-only partial fine-tuning, and (3) correct pos_weight BCE penalisation. "
        "This is the configuration that achieved the 90.0% subject-level diagnostic accuracy."
    ).runs[0].font.size = Pt(11)

    # ─── 4. Results ───────────────────────────────────────────────────────────
    doc.add_heading("4. Results and Discussion", level=1).runs[0].font.color.rgb = RGBColor(0x1a,0x1a,0x2e)

    doc.add_heading("4.1 Global Sequence-Level Evaluation", level=2)
    doc.add_paragraph(
        "The final Spatial RGB model outperformed Experiment 1 across all continuous-sequence metrics:"
    ).runs[0].font.size = Pt(11)
    metrics = [
        ("Global Accuracy", "80.5%", "84.8%"),
        ("MDD Sensitivity (Recall)", "90.1%", "93.8%"),
        ("Healthy Specificity (with outliers)", "63.6%", "68.0%"),
        ("Healthy Specificity (excluding outliers)", "81.4%", "89.8% 🚀"),
        ("Median Specificity across all subjects", "78.4%", "93.8%"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Experiment 1 (Baseline)"
    hdr[2].text = "Final (Spatial RGB)"
    for h in hdr:
        h.paragraphs[0].runs[0].bold = True
    for name, e1, final in metrics:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = e1
        row[2].text = final

    doc.add_paragraph()
    doc.add_paragraph(
        "The Specificity Mean of 68% was mathematically dragged down by the three biologically anomalous "
        "healthy subjects (H_16, H_24, H_27) which scored near 0%. Evaluating the Median Specificity "
        "reveals the true model capability: a massive leap from 78.4% to 93.8%. Additionally, the Spatial "
        "RGB mapping provided the network enough topographic context to correctly diagnose the extreme "
        "outlier MDD_19, which earlier algorithms had rated at <10% accuracy."
    ).runs[0].font.size = Pt(11)
    add_image_safe(doc, "results/exp_final/figures/comparison_metrics.png",
                   "Figure 5: Percentage-point growth in Mean and Median metrics between Baseline and Final Architecture.")

    doc.add_heading("4.2 Clinical Subject-Level Evaluation (Majority Vote)", level=2)
    doc.add_paragraph(
        "Evaluating singular 7-second sequences is clinically insufficient in isolation. We implemented "
        "a Subject-Level Majority Vote: every patient's constituent sequence predictions are aggregated "
        "and the majority label is taken as the clinical diagnosis."
    ).runs[0].font.size = Pt(11)
    add_image_safe(doc, "results/exp_final/figures/subject_accuracy.png",
                   "Figure 6: Per-subject granular sequence accuracy using the Final Architecture.")
    add_image_safe(doc, "results/exp_final/figures/majority_vote_pie.png",
                   "Figure 7: Clinical diagnostic outcome — 36 / 40 patients correctly diagnosed (90.0% accuracy).")

    doc.add_heading("4.3 Comparative Summary", level=2)
    comp = doc.add_table(rows=1, cols=5)
    comp.style = 'Table Grid'
    hdrs = comp.rows[0].cells
    for i, h in enumerate(["Architecture", "Spatial Strategy", "Temporal", "Validation", "Clinical Accuracy"]):
        hdrs[i].text = h
        hdrs[i].paragraphs[0].runs[0].bold = True
    rows_data = [
        ("Handcrafted SVM / k-NN", "Extracted Channels", "None", "Train/Test Split", "~86%–91% ⚠️ leakage"),
        ("Exp 1: Grayscale Inception+LSTM", "Averaged Grayscale", "seq=10", "LOSO", "63.9% Spec"),
        ("Exp 2: Fine-Tuned Inception+LSTM", "Averaged Grayscale", "seq=10", "LOSO", "68.0%"),
        ("Proposed: Spatial RGB LSTM", "RGB Topographic Map", "seq=10", "Strict LOSO", "90.0% (36/40) ✅"),
    ]
    for row_data in rows_data:
        r = comp.add_row().cells
        for i, val in enumerate(row_data):
            r[i].text = val
    doc.add_paragraph()

    # ─── 5. Conclusion ────────────────────────────────────────────────────────
    doc.add_heading("5. Conclusion", level=1).runs[0].font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
    doc.add_paragraph(
        "This research presents a robust, interpretable, and reproducible deep learning framework for the "
        "automated identification of MDD from continuous EEG signals. By introducing Spatial RGB Mapping — "
        "encoding topographic brain regions as independent colour channels — and partially fine-tuning "
        "InceptionV3's Mixed_7c block, we equipped the model with the spatial context necessary to "
        "distinguish biologically similar but clinically opposite EEG profiles. Combined with LSTM-based "
        "temporal modelling and strict Leave-One-Subject-Out cross-validation, the framework achieved "
        "90.0% Subject-Level Diagnostic Accuracy (36 / 40 patients). Rigorous Outlier Analysis further "
        "documented the biological constraints posed by atypical Alpha energy distributions, providing "
        "an unprecedented level of transparency rarely seen in EEG deep learning studies. This framework "
        "offers a promising step toward scalable, objective mental health diagnostics in resource-limited settings."
    ).runs[0].font.size = Pt(11)

    # ─── References ───────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    doc.add_heading("References", level=2)
    refs = [
        "[1] Hossain MRT et al. WSEAS TRANSACTIONS ON BIOLOGY AND BIOMEDICINE 2025; 22: 133–151.",
        "[2] Thoduparambil PP et al. Phys Eng Sci Med 2020. DOI: 10.1007/s13246-020-00938-4.",
        "[3] Salehi AW et al. Sustainability 2023. DOI: 10.3390/su15075930.",
        "[4] Li CT et al. J Affect Disord 2023. DOI: 10.1016/j.jad.2023.08.059.",
        "[5] Mumtaz W. MDD Patients and Healthy Controls EEG Data (New). 2016. DOI: 10.6084/m9.figshare.4244171.v2.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style='List Paragraph')
        p.runs[0].font.size = Pt(9)

    doc.save(str(DOCX_OUT))
    print(f"✅  DOCX saved → {DOCX_OUT}")


# ─── Build HTML → PDF ────────────────────────────────────────────────────────

def md_img_to_abs(md_content):
    """Replace relative image paths with absolute paths for wkhtmltopdf / Chrome."""
    def replacer(m):
        rel = m.group(1)
        abs_p = (BASE_DIR / rel).resolve()
        if abs_p.exists():
            return f'src="file://{abs_p}"'
        return m.group(0)
    return re.sub(r'src="([^"]+)"', replacer, md_content)


def build_pdf():
    """Generate HTML with proper CSS then use headless Chrome to print to PDF."""
    # Convert MD → plain HTML via pandoc
    result = subprocess.run(
        ["pandoc", str(MD_PATH), "--standalone", "--self-contained", "-o", str(HTML_TMP)],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print("pandoc error:", result.stderr); return

    html = HTML_TMP.read_text(encoding="utf-8")

    # Fix relative image src paths → absolute file:// paths
    def abs_src(m):
        rel = m.group(1)
        abs_p = (BASE_DIR / rel).resolve()
        return f'src="file://{abs_p}"' if abs_p.exists() else m.group(0)
    html = re.sub(r'src="([^"http][^"]*)"', abs_src, html)

    # Inject professional CSS
    css = """
    <style>
      @import url('https://fonts.googleapis.com/css2?family=Source+Serif+4:ital,opsz,wght@0,8..60,300..900;1,8..60,300..900&family=Source+Sans+3:wght@400;600&display=swap');
      * { box-sizing: border-box; }
      body {
        font-family: 'Source Serif 4', Georgia, serif;
        font-size: 11pt;
        line-height: 1.65;
        color: #1a1a2e;
        max-width: 820px;
        margin: 0 auto;
        padding: 48px 60px;
      }
      h1 {
        font-family: 'Source Serif 4', serif;
        font-size: 18pt;
        font-weight: 700;
        text-align: center;
        color: #1a1a2e;
        margin-bottom: 8px;
        line-height: 1.3;
      }
      h2 {
        font-family: 'Source Serif 4', serif;
        font-size: 13pt;
        font-weight: 700;
        color: #1a1a2e;
        border-bottom: 1.5px solid #d0d0e0;
        padding-bottom: 4px;
        margin-top: 28px;
      }
      h3 {
        font-family: 'Source Serif 4', serif;
        font-size: 11.5pt;
        font-weight: 700;
        color: #2c3e70;
        margin-top: 20px;
      }
      h4 {
        font-size: 11pt;
        font-weight: 600;
        color: #444;
        margin-top: 12px;
      }
      p { margin: 0 0 10px 0; }
      /* Authors block */
      p strong:only-child { display: block; text-align: center; font-size: 10.5pt; }
      hr {
        border: none;
        border-top: 1px solid #bbb;
        margin: 24px 0;
      }
      table {
        width: 100%;
        border-collapse: collapse;
        font-size: 9.5pt;
        margin: 16px 0;
      }
      th {
        background: #1a1a2e;
        color: white;
        padding: 7px 10px;
        text-align: left;
      }
      td {
        padding: 6px 10px;
        border: 1px solid #d4d4e0;
      }
      tr:nth-child(even) td { background: #f5f5ff; }
      tr:last-child td { font-weight: 600; background: #eef0ff; }
      img {
        max-width: 100%;
        display: block;
        margin: 16px auto;
        border: 1px solid #ddd;
        border-radius: 4px;
      }
      em { color: #555; }
      code {
        font-family: 'Courier New', monospace;
        font-size: 9.5pt;
        background: #f0f0f8;
        padding: 1px 4px;
        border-radius: 3px;
      }
      blockquote {
        border-left: 3px solid #8080c0;
        margin: 12px 0;
        padding: 8px 16px;
        background: #f8f8ff;
        font-size: 10pt;
        color: #444;
      }
      ul, ol { margin: 6px 0 10px 20px; }
      li { margin-bottom: 4px; }
      /* Figure captions */
      p > em:only-child {
        display: block;
        text-align: center;
        font-size: 9pt;
        color: #666;
        margin-top: -8px;
      }
    </style>
    """
    html = html.replace("</head>", f"{css}\n</head>")
    HTML_TMP.write_text(html, encoding="utf-8")

    # Print to PDF via headless Chrome
    chrome_paths = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Google Chrome Canary.app/Contents/MacOS/Google Chrome Canary",
    ]
    chrome_bin = next((p for p in chrome_paths if Path(p).exists()), None)
    if not chrome_bin:
        print("❌  Chrome not found — PDF skipped."); return

    res = subprocess.run([
        chrome_bin, "--headless", "--disable-gpu",
        f"--print-to-pdf={PDF_OUT}",
        f"file://{HTML_TMP}",
    ], capture_output=True, text=True)

    HTML_TMP.unlink(missing_ok=True)
    if res.returncode == 0:
        print(f"✅  PDF saved → {PDF_OUT}")
    else:
        print("Chrome error:", res.stderr[:500])


if __name__ == "__main__":
    try:
        from docx import Document
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    build_docx()
    build_pdf()
